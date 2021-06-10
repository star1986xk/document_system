import os
import sys
import pythoncom
from win32com import client as wc
from docx import Document
import openpyxl
from PyQt5.QtCore import Qt
from PyQt5.QtWidgets import QApplication, QFrame, QMessageBox, QFileDialog
from UI.ui_main import Ui_Form
from login_win import LoginWin


class mainMain(Ui_Form, QFrame):
    def __init__(self, parent=None):
        '''
        程序初始化
        :param parent:
        '''
        super().__init__(parent)
        self.setupUi(self)
        self.retranslateUi(self)

        self.setWindowFlags(
            Qt.Window | Qt.FramelessWindowHint | Qt.WindowSystemMenuHint | Qt.WindowMinimizeButtonHint | Qt.WindowMaximizeButtonHint)  # 窗口无边框

        self.login = LoginWin()
        self.login.Sig_login.connect(self.load)
        self.login.show()

        self.toolButton_word.clicked.connect(self.open_word)
        self.toolButton_excel.clicked.connect(self.out_excel)
        self.toolButton_run.clicked.connect(self.run)
        self.toolButton_close.clicked.connect(self.close)

    # doc转成docx
    def toDOCX(self, filename):
        try:
            pythoncom.CoInitialize()
            word = wc.Dispatch('Word.Application')
            doc = word.Documents.Open(filename)
            doc.SaveAs(filename + 'x', 12)
            doc.Close()
            word.Quit()
            os.remove(filename)
            pythoncom.CoUninitialize()
        except Exception as e:
            pass
        finally:
            return filename + 'x'

    # 解析word表格
    def readWord(self, path):
        try:
            document = Document(path)  # 读入文件
            tables = document.tables  # 获取文件中的表格集

            jiandang = document.paragraphs[1].text.replace('建档时间：', '')
            xingming = tables[0].rows[0].cells[1].text
            cengyongming = tables[0].rows[1].cells[1].text
            xingbie = tables[0].rows[2].cells[1].text
            minzu = tables[0].rows[2].cells[4].text
            jiguan = tables[0].rows[3].cells[1].text
            chushengdi = tables[0].rows[3].cells[4].text
            mianmao = tables[0].rows[4].cells[1].text
            xinyang = tables[0].rows[4].cells[4].text
            gonghao = tables[0].rows[5].cells[1].text
            xueli = tables[0].rows[5].cells[4].text
            danwei = tables[0].rows[6].cells[1].text
            zhiwu = tables[0].rows[7].cells[1].text
            zhicheng = tables[0].rows[7].cells[4].text
            shenfenzheng = tables[0].rows[8].cells[4].text
            sfz_dizhi = tables[0].rows[9].cells[4].text
            jiating_dizhi = tables[0].rows[10].cells[4].text
            changzhu_dizhi = tables[0].rows[11].cells[4].text
            shouji = tables[0].rows[12].cells[2].text
            qq = tables[0].rows[12].cells[6].text
            weixin = tables[0].rows[13].cells[2].text
            youxiang = tables[0].rows[13].cells[6].text
            biyeshijian = tables[0].rows[14].cells[4].text
            rudang = tables[0].rows[15].cells[4].text
            rutuan = tables[0].rows[16].cells[4].text

            p_xingming = tables[1].rows[0].cells[4].text
            p_xingbie = tables[1].rows[0].cells[8].text
            p_shengri = tables[1].rows[1].cells[4].text
            p_chushengdi = tables[1].rows[1].cells[8].text
            p_minzu = tables[1].rows[2].cells[4].text
            p_jiguan = tables[1].rows[2].cells[8].text
            p_shenfenzheng = tables[1].rows[3].cells[6].text
            p_zhengzhi = tables[1].rows[4].cells[4].text
            p_xinyang = tables[1].rows[4].cells[8].text
            p_shouji = tables[1].rows[5].cells[4].text
            p_qq = tables[1].rows[5].cells[8].text
            p_weixin = tables[1].rows[6].cells[4].text
            p_youxiang = tables[1].rows[6].cells[8].text
            p_zhiwu = tables[1].rows[7].cells[4].text

            junxian = tables[2].rows[0].cells[1].text
            chuguo = tables[2].rows[1].cells[1].text
            huojiang = tables[2].rows[2].cells[1].text
            chufen = tables[2].rows[3].cells[1].text
            fandong = tables[2].rows[4].cells[1].text

            g_shijian = tables[3].rows[1].cells[1].text
            g_didian = tables[3].rows[2].cells[1].text
            g_juese = tables[3].rows[3].cells[1].text
            g_qita = tables[3].rows[4].cells[1].text
            g_qingkuang = tables[3].rows[5].cells[1].text

            return [jiandang, xingming, cengyongming, xingbie, minzu, jiguan, chushengdi, mianmao, xinyang, gonghao, xueli,
                    danwei, zhiwu, zhicheng, shenfenzheng, sfz_dizhi, jiating_dizhi, changzhu_dizhi, shouji, qq, weixin,
                    youxiang, biyeshijian, rudang, rutuan, p_xingming, p_xingbie, p_shengri, p_chushengdi, p_minzu,
                    p_jiguan, p_shenfenzheng, p_zhengzhi, p_xinyang, p_shouji, p_qq, p_weixin, p_youxiang, p_zhiwu, junxian,
                    chuguo, huojiang, chufen, fandong, g_shijian, g_didian, g_juese, g_qita, g_qingkuang]
        except Exception as e:
            print(e)

    # 保存excel
    def save(self, path, data):
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.append(
            ['建档时间', '姓名', '曾用名', '性别', '名族', '籍贯', '出生地', '政治面貌', '宗教信仰', '工号/学号', '学历/学位', '单位/学院', '职务', '职称',
             '身份证号码', '身份证地址', '户籍/家庭地址', '常住地址/学校住址', '手机/座机', 'QQ', '微信', '电子邮箱', '何年何月何处参加工作入学/毕业时间',
             '何年何月何人介绍加入中国共产党，何时转正', '何年何月加入中国共产主义青年团', '配偶姓名', '配偶性别', '配偶出生年月', '配偶出生地', '配偶民族', '配偶籍贯', '配偶身份证号',
             '配偶政治面貌', '配偶信教信仰', '配偶手机/座机', '配偶QQ', '配偶微信', '配偶电子邮箱', '配偶工作单位及职务', '获得军、警衔', '出国情况', '获奖情况', '受处分情况',
             '参加反动组织，任何职务，有何结论情况', '关注事项记录--起止时间', '关注事项记录--地点', '关注事项记录--扮演角色', '关注事项记录--其它人员', '关注事项记录--情况描述'])
        for li in data:
            ws.append(li)
        wb.save(path)

    # 登陆成功显示窗口
    def load(self):
        self.show()

    # 选择导入文件
    def open_word(self):
        try:
            filenames, _ = QFileDialog.getOpenFileNames(self, "选取Word文件", ".", "Word格式 (*.doc; *.docx)")
            self.lineEdit_word.setText(','.join(filenames))
        except Exception as e:
            print(e)

    # 选择导出位置
    def out_excel(self):
        try:
            filenames, _ = QFileDialog.getSaveFileName(self, '保存文件', "导出.xlsx", "Excel格式 (*.xlsx)")
            self.lineEdit_excel.setText(filenames)
        except Exception as e:
            print(e)

    # 运行
    def run(self):
        word_path = self.lineEdit_word.text()
        excel_path = self.lineEdit_excel.text()
        if not word_path or not excel_path: return
        data = []
        for path in word_path.split(','):
            name, ext = os.path.splitext(path)
            if ext.lower() == '.doc':
                filename = self.toDOCX(path)
            else:
                filename = path
            data.append(self.readWord(filename))
        self.save(excel_path, data)
        QMessageBox.information(self, '提示', '导出成功')

    # 重写窗口鼠标拖动事件
    def mousePressEvent(self, event):
        if event.button() == Qt.LeftButton:
            self.m_flag = True
            self.m_Position = event.globalPos() - self.pos()  # 获取鼠标相对窗口的位置
            event.accept()

    def mouseMoveEvent(self, QMouseEvent):
        if Qt.LeftButton and self.m_flag:
            self.move(QMouseEvent.globalPos() - self.m_Position)  # 更改窗口位置
            QMouseEvent.accept()

    def mouseReleaseEvent(self, QMouseEvent):
        self.m_flag = False


if __name__ == '__main__':
    app = QApplication(sys.argv)
    win = mainMain()
    # win.show()
    sys.exit(app.exec_())
